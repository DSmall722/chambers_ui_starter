
import os, json
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

BASE_DIR = os.path.dirname(__file__)
DATA_DIR = os.path.join(BASE_DIR, "data")
os.makedirs(DATA_DIR, exist_ok=True)

def load_json(name: str, default):
    path = os.path.join(DATA_DIR, name)
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default

def save_json(name: str, obj):
    path = os.path.join(DATA_DIR, name)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

APP_MEMOS = os.path.join(DATA_DIR, "memos")
APP_EXPORTS = os.path.join(DATA_DIR, "exports")
APP_UPLOADS = os.path.join(DATA_DIR, "uploads")
for d in (APP_MEMOS, APP_EXPORTS, APP_UPLOADS):
    os.makedirs(d, exist_ok=True)

app.mount('/memos', StaticFiles(directory=APP_MEMOS), name='memos')
app.mount('/exports', StaticFiles(directory=APP_EXPORTS), name='exports')
app.mount('/files', StaticFiles(directory=APP_UPLOADS), name='files')

# ---------- Models ----------
class ChatTurn(BaseModel):
    role: str
    content: str
    matter_id: Optional[str] = None

class ChatRequest(BaseModel):
    matter_id: Optional[str] = None
    message: str
    history: List[ChatTurn] = []
    personas: Dict[str, str] = {}
    dials: Dict[str, Any] = {}

class MemoSave(BaseModel):
    matter_id: Optional[str] = None
    title: Optional[str] = None
    history: List[Dict[str, Any]]

class DrafterReq(BaseModel):
    matter_id: Optional[str] = None
    goal: Optional[str] = None
    outline: List[str] = []
    citations: List[str] = []
    persona: Optional[str] = None
    dials: Dict[str, Any] = {}

class EditReq(BaseModel):
    text: str
    constraints: Dict[str, Any] = {}

class CiteReq(BaseModel):
    text: str
    forum: Optional[str] = None

class RulesReq(BaseModel):
    forum: Optional[str] = None
    doc_type: Optional[str] = None
    content: Optional[str] = None

class RAGReq(BaseModel):
    matter_id: Optional[str] = None
    query: str
    k: int = 5

class SchedReq(BaseModel):
    forum: str
    trigger: Dict[str, Any]

class Matter(BaseModel):
    id: str
    title: str
    forum: Optional[str] = None
    client: Optional[str] = None

class Contact(BaseModel):
    id: str
    name: str
    roles: List[str] = []
    org: Optional[str] = None
    emails: List[str] = []
    matters: List[str] = []
    tags: List[str] = []
    note: Optional[str] = None
    last_activity: Optional[int] = None

# ---------- Activity ----------
def _activity_path():
    return os.path.join(DATA_DIR, "activity.json")

def load_activity():
    if not os.path.exists(_activity_path()):
        return {}
    try:
        return json.load(open(_activity_path(), "r", encoding="utf-8"))
    except Exception:
        return {}

def save_activity(obj):
    json.dump(obj, open(_activity_path(), "w", encoding="utf-8"), indent=2)

def touch_activity(matter_id: str, kind: str):
    if not matter_id:
        return
    act = load_activity()
    now = int(datetime.utcnow().timestamp())
    rec = act.get(matter_id) or {"last_activity": 0, "counts": {}}
    rec["last_activity"] = now
    rec["counts"][kind] = int(rec["counts"].get(kind, 0)) + 1
    act[matter_id] = rec
    save_activity(act)

# ---------- Endpoints ----------
@app.post("/chat")
def chat(req: ChatRequest):
    # Stub: echo back with persona/dial info
    reply = f"Understood. Personas: {req.personas}. Dials: {req.dials}."
    out = {"reply": reply, "needs": []}
    # Example: if user says "draft", return a 'draft' result for UI export
    if "draft" in (req.message or "").lower():
        out["result"] = {"type":"draft","persona": req.personas.get("drafter","Aquila (Drafter)"), "text": "DRAFT — placeholder text based on your outline and citations."}
    touch_activity(req.matter_id or "", "chat")
    return out

@app.post("/chat/save_memo")
def save_memo(req: MemoSave):
    when = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    name = f"{req.matter_id or 'general'}-{when}.md"
    path = os.path.join(APP_MEMOS, name)
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"# Memo: {req.title or ''}\n\n")
        for turn in req.history:
            role = turn.get("role","?")
            content = turn.get("content","")
            f.write(f"**{role}**\n\n{content}\n\n---\n\n")
    return {"ok": True, "download_url": f"/memos/{name}"}

@app.post("/preset/suggest")
def preset_suggest(payload: Dict[str, Any]):
    text = (payload or {}).get("complaint_text","").lower()
    if "removal" in text or "federal" in text:
        return {"suggested_preset":"Removal Pack"}
    return {"suggested_preset":"State Answer — Standard Defense"}

@app.post("/initial_review")
def initial_review(payload: Dict[str, Any]):
    files = []
    mid = (payload or {}).get("matter_id","")
    base = os.path.join(APP_UPLOADS, mid)
    if os.path.isdir(base):
        files = sorted([f for f in os.listdir(base) if not f.startswith(".")])
    checklist = ["Summons present","Complaint present","Service verified","Cover sheet filled"]
    touch_activity(mid, "initial_review")
    return {"files": files, "checklist": checklist}

@app.post("/drafter")
def drafter(req: DrafterReq):
    text = f"# {req.goal or 'Draft'}\n\n" + "\n".join([f"## {h}" for h in (req.outline or ["Introduction","Argument","Conclusion"])])
    return {"persona": req.persona, "draft": text, "dials": req.dials}

@app.post("/edit")
def edit(req: EditReq):
    txt = (req.text or "").strip()
    max_cut = float(req.constraints.get("max_cut", 0.2))
    # trivial "edit": collapse spaces, cut to ~ (1-max_cut) chars
    cleaned = " ".join(txt.split())
    target_len = max(int(len(cleaned) * (1.0 - max_cut)), 0)
    edited = cleaned[:target_len]
    return {"edited": edited}

@app.post("/cite/lint")
def cite_lint(req: CiteReq):
    return {"ok": True, "forum": req.forum or "SC:Trial", "issues": []}

@app.post("/rules/validate")
def rules_validate(req: RulesReq):
    return {"ok": True, "forum": req.forum or "SC:Trial", "doc_type": req.doc_type or "brief", "issues": []}

@app.post("/rag/search")
def rag_search(req: RAGReq):
    # simple echo stub
    return {"results": [{"title":"Example doc","snippet":"...example...","score":0.42}]}

@app.post("/schedule/compute")
def schedule_compute(req: SchedReq):
    dt = req.trigger.get("datetime") or datetime.utcnow().isoformat()
    # add 30 days as a demo
    due = (datetime.fromisoformat(dt.replace("Z","")) + timedelta(days=30)).isoformat()
    return {"final_due": due}

@app.get("/matters", response_model=List[Matter])
def matters():
    return load_json("matters.json", [])

@app.post("/matters")
def add_matter(payload: Dict[str, Any]):
    items = load_json("matters.json", [])
    # naive id if not provided
    if "id" not in payload or not payload["id"]:
        base = (payload.get("title","MAT")[:3] or "MAT").upper()
        num = len(items) + 1
        payload["id"] = f"{base}{num:03d}"
    items.append(payload)
    save_json("matters.json", items)
    touch_activity(payload["id"], "matter_add")
    return payload

# Contacts
@app.get("/contacts", response_model=List[Contact])
def contacts(q: Optional[str] = None, role: Optional[str] = None, matter_id: Optional[str] = None, has_email: Optional[int] = None):
    data = load_json("contacts.json", [])
    if not any([q, role, matter_id, has_email]):
        return data
    qlow = (q or "").strip().lower()
    out = []
    for c in data:
        if role and role not in (c.get("roles") or []):
            continue
        if matter_id and matter_id not in (c.get("matters") or []):
            continue
        if has_email and not (c.get("emails") or []):
            continue
        if qlow:
            hay = " ".join([
                c.get("name",""),
                " ".join(c.get("roles") or []),
                c.get("org","") or "",
                " ".join(c.get("emails") or []),
                " ".join(c.get("matters") or []),
                " ".join(c.get("tags") or []),
                c.get("note","") or ""
            ]).lower()
            if qlow not in hay:
                continue
        out.append(c)
    return out

class LinkContact(BaseModel):
    id: str
    matter_id: str

@app.post("/contacts/link")
def link_contact(req: LinkContact):
    contacts = load_json("contacts.json", [])
    for c in contacts:
        if c.get("id") == req.id:
            ms = c.get("matters") or []
            if req.matter_id not in ms:
                ms.append(req.matter_id)
                c["matters"] = ms
            save_json("contacts.json", contacts)
            return {"ok": True, "contact": c}
    return {"ok": False, "error": "Not found"}

@app.post("/contacts")
def add_contact(payload: Dict[str, Any]):
    contacts = load_json("contacts.json", [])
    if "id" not in payload or not payload["id"]:
        base = (payload.get("name","CON")[:3] or "CON").upper()
        num = len(contacts) + 1
        payload["id"] = f"{base}{num:04d}"
    contacts.append(payload)
    save_json("contacts.json", contacts)
    return payload

@app.delete("/contacts")
def delete_contact(id: str):
    contacts = load_json("contacts.json", [])
    newc = [c for c in contacts if c.get("id") != id]
    if len(newc) == len(contacts):
        return {"ok": False, "error": "Not found"}
    save_json("contacts.json", newc)
    return {"ok": True}

# Uploads
def safe_name(name: str) -> str:
    keep = [c for c in name if c.isalnum() or c in (' ','-','_','.',')','(')]
    return ''.join(keep).strip().replace('..','.')

@app.post("/upload")
async def upload_file(matter_id: str, file: UploadFile = File(...)):
    name = file.filename or "upload"
    kind = "exhibit"
    low = name.lower()
    if low.endswith(".pdf"):
        if "order" in low:
            kind = "order"
        elif "brief" in low or "memo" in low:
            kind = "brief"
    upbase = os.path.join(APP_UPLOADS, matter_id)
    os.makedirs(upbase, exist_ok=True)
    when = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    safe = safe_name(name)
    final = f"{when}-{safe}"
    path = os.path.join(upbase, final)
    content = await file.read()
    with open(path, "wb") as f:
        f.write(content)
    index_path = os.path.join(upbase, ".index.json")
    try:
        meta = json.load(open(index_path,"r",encoding="utf-8"))
    except Exception:
        meta = {}
    meta.setdefault(final, {"tags": [], "note": ""})
    json.dump(meta, open(index_path,"w",encoding="utf-8"), indent=2)
    touch_activity(matter_id, "upload")
    return {"ok": True, "matter_id": matter_id, "filename": final, "original": name, "detected": kind}

@app.get("/uploads")
def list_uploads(matter_id: str):
    base = os.path.join(APP_UPLOADS, matter_id)
    files = []
    meta = {}
    index_path = os.path.join(base, ".index.json")
    if os.path.exists(index_path):
        try:
            meta = json.load(open(index_path,"r",encoding="utf-8"))
        except Exception:
            meta = {}
    if os.path.isdir(base):
        for fn in sorted(os.listdir(base)):
            if fn.startswith("."):
                continue
            p = os.path.join(base, fn)
            if os.path.isfile(p):
                st = os.stat(p)
                m = meta.get(fn, {"tags": [], "note": ""})
                files.append({"filename": fn, "size": st.st_size, "modified": st.st_mtime, "tags": m.get("tags",[]), "note": m.get("note","")})
    return {"matter_id": matter_id, "files": files}

@app.delete("/uploads")
def delete_upload(matter_id: str, filename: str):
    base = os.path.join(APP_UPLOADS, matter_id)
    path = os.path.join(base, filename)
    if os.path.isfile(path):
        os.remove(path)
        index_path = os.path.join(base, ".index.json")
        try:
            meta = json.load(open(index_path,"r",encoding="utf-8"))
        except Exception:
            meta = {}
        if filename in meta:
            meta.pop(filename, None)
            json.dump(meta, open(index_path,"w",encoding="utf-8"), indent=2)
        return {"ok": True}
    return {"ok": False, "error": "Not found"}

class UploadMeta(BaseModel):
    matter_id: str
    filename: str
    tags: List[str] = []
    note: str = ""

@app.post("/uploads/meta")
def set_upload_meta(req: UploadMeta):
    base = os.path.join(APP_UPLOADS, req.matter_id)
    os.makedirs(base, exist_ok=True)
    index_path = os.path.join(base, ".index.json")
    try:
        meta = json.load(open(index_path,"r",encoding="utf-8"))
    except Exception:
        meta = {}
    meta[req.filename] = {"tags": req.tags or [], "note": req.note or ""}
    json.dump(meta, open(index_path,"w",encoding="utf-8"), indent=2)
    touch_activity(req.matter_id, "meta")
    return {"ok": True}

# Recent matters
@app.get("/matters_recent")
def matters_recent():
    matters = load_json("matters.json", [])
    act = load_activity()
    out = []
    for m in matters:
        aid = m["id"]
        meta = act.get(aid, {"last_activity":0,"counts":{}})
        out.append({
            "id": aid,
            "title": m.get("title",""),
            "forum": m.get("forum",""),
            "client": m.get("client",""),
            "last_activity": meta.get("last_activity",0),
            "counts": meta.get("counts",{})
        })
    out.sort(key=lambda x: x.get("last_activity",0), reverse=True)
    return out

# Export DOCX
class ExportDocxRequest(BaseModel):
    text: str
    filename: Optional[str] = None
    title: Optional[str] = None

@app.post("/export/docx")
def export_docx(req: ExportDocxRequest):
    try:
        from docx import Document
    except Exception as e:
        return {"ok": False, "error": "python-docx not installed", "detail": str(e)}
    when = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    safe = (req.filename or f"draft-{when}").strip().replace("/", "-")[:80]
    if not safe.lower().endswith(".docx"):
        safe += ".docx"
    path = os.path.join(APP_EXPORTS, safe)
    doc = Document()
    if req.title:
        doc.add_heading(req.title, level=1)
    blocks = req.text.replace("\r\n","\n").split("\n\n")
    for b in blocks:
        for line in b.split("\n"):
            doc.add_paragraph(line)
        doc.add_paragraph()
    doc.save(path)
    return {"ok": True, "filename": safe, "download_url": f"/exports/{safe}"}


# ---- Multi-persona Compare (API) ----
class MultiRunReq(BaseModel):
    tool: str  # "initial" | "drafter" | "editor"
    personas: List[str] = []
    payload: Dict[str, Any] = {}
    rounds: int = 1
    matter_id: Optional[str] = None
    debate: Optional[bool] = False

class MultiRunOut(BaseModel):
    runs: Dict[str, Dict[str, Any]]
    synthesis: Dict[str, Any]
    trace_id: str

def _mp_synthesize(runs: Dict[str, Dict[str, Any]], tool: str, personas: List[str], debate: bool) -> Dict[str, Any]:
    # Build a simple synthesis from persona texts.
    parts = [
        f"**Synthesis — {tool}**",
        f"Personas: {', '.join(personas)}",
        f"Debate mode: {'ON' if debate else 'OFF'}",
        ""
    ]
    seen = set()
    for name, r in runs.items():
        text = (r.get("text") or "").splitlines()
        for line in text:
            s = line.strip()
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            parts.append(f"- {s}")
    return {"text": "\\n".join(parts), "sources": []}

def _mp_markdown(matter_id: str, tool: str, personas: List[str], synthesis: Dict[str, Any], runs: Dict[str, Dict[str, Any]], rounds: int, debate: bool) -> str:
    from datetime import datetime
    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        f"# Multi-Persona Synthesis — {matter_id}",
        f"*Tool:* {tool}  ",
        f"*Personas:* {', '.join(personas)}  ",
        f"*Rounds:* {rounds}  *Debate:* {'ON' if debate else 'OFF'}  ",
        f"*Generated:* {ts}",
        "",
        "## Synthesis (Concordia)",
        synthesis.get("text",""),
        ""
    ]
    for p, r in runs.items():
        lines += [f"## {p}", r.get("text",""), ""]
        srcs = r.get("sources") or []
        if srcs:
            lines += ["**Sources**"] + [f"- {s}" for s in srcs] + [""]
    return "\\n".join(lines)

@app.post("/tool/multi_run")
def tool_multi_run(req: MultiRunReq):
    # Placeholder execution for each persona (swap with real model calls later).
    runs: Dict[str, Dict[str, Any]] = {}
    for persona in req.personas:
        summary = [f"{persona} — {req.tool} run:"]
        keys = list(req.payload.keys())[:8]
        if keys:
            summary.append("Inputs: " + ", ".join(f"{k}={str(req.payload[k])[:40]}" for k in keys))
        if req.debate and req.rounds and req.rounds >= 2:
            summary.append("Round 2: critique phase (placeholder)")
        runs[persona] = {"text": "\\n".join(summary), "sources": []}

    synthesis = _mp_synthesize(runs, req.tool, req.personas, bool(req.debate))
    trace_id = __import__("datetime").datetime.utcnow().strftime("mp-%Y%m%d-%H%M%S")

    # Auto-save memo
    try:
        md = _mp_markdown(req.matter_id or "UNKNOWN", req.tool, req.personas, synthesis, runs, req.rounds or 1, bool(req.debate))
        import os
        memdir = os.path.join(DATA_DIR, "memos")
        os.makedirs(memdir, exist_ok=True)
        safe = (req.matter_id or "UNKNOWN").replace("/", "_")
        fname = f"multi-{safe}-{trace_id}.md"
        with open(os.path.join(memdir, fname), "w", encoding="utf-8") as f:
            f.write(md)
    except Exception:
        pass

    return {"runs": runs, "synthesis": synthesis, "trace_id": trace_id}

class MultiSaveMemo(BaseModel):
    matter_id: str
    title: str
    markdown: str

@app.post("/multi/save_memo")
def multi_save_memo(req: MultiSaveMemo):
    import os
    memdir = os.path.join(DATA_DIR, "memos")
    os.makedirs(memdir, exist_ok=True)
    safe = req.matter_id.replace("/", "_")
    ts = __import__("datetime").datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    fname = f"{req.title.replace(' ', '_')}-{safe}-{ts}.md"
    path = os.path.join(memdir, fname)
    with open(path, "w", encoding="utf-8") as f:
        f.write(req.markdown)
    return {"ok": True, "download_url": f"/memos/{fname}"}
